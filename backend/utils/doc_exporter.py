import re
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT_COLOR = RGBColor(0x1F, 0x3C, 0x5E)   # deep navy -- reads as professional, prints fine in B/W
BODY_COLOR = RGBColor(0x22, 0x22, 0x22)
MUTED_COLOR = RGBColor(0x55, 0x55, 0x55)

FONT_NAME = "Calibri"  # universally available, ATS-safe, default Word font

_INLINE_RE = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")


def _add_inline_runs(paragraph, text, base_size=Pt(10.5), color=BODY_COLOR):
    """Render **bold** / *italic* markdown spans as real bold/italic runs
    instead of leaking literal asterisks into the document."""
    for token in _INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(token)
        run.font.name = FONT_NAME
        run.font.size = base_size
        run.font.color.rgb = color


def _set_bottom_border(paragraph, color="1F3C5E", size=6):
    """Add a single bottom border to a paragraph -- used under section
    headings instead of a separate horizontal-rule element."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_borders.append(bottom)
    p_pr.append(p_borders)


def _add_bullet(doc, text, indent=0.25):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    _add_inline_runs(p, text)
    return p


def _markdown_to_resume_docx(doc, lines):
    i = 0
    name_done = False

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            # Candidate name -- large, centered, accent color
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:].strip())
            run.font.name = FONT_NAME
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = ACCENT_COLOR
            p.paragraph_format.space_after = Pt(2)
            name_done = True
            i += 1
            continue

        if name_done and not stripped.startswith(("#", "-", "*", "**")):
            # The line right after the name is the contact line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.font.name = FONT_NAME
            run.font.size = Pt(9.5)
            run.font.color.rgb = MUTED_COLOR
            p.paragraph_format.space_after = Pt(10)
            name_done = False
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[3:].strip().upper())
            run.font.name = FONT_NAME
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = ACCENT_COLOR
            _set_bottom_border(p)
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            _add_bullet(doc, stripped[2:].strip())
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            # Role/Company/Dates line e.g. **Senior Engineer | Acme | 2021-Present**
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            _add_inline_runs(p, stripped, base_size=Pt(10.5))
            for run in p.runs:
                run.bold = True
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_inline_runs(p, stripped)
        i += 1


def _markdown_to_cover_letter_docx(doc, lines):
    name_done = False
    for raw in lines:
        stripped = raw.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("# ") and not name_done:
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:].strip())
            run.font.name = FONT_NAME
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = ACCENT_COLOR
            p.paragraph_format.space_after = Pt(10)
            name_done = True
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        _add_inline_runs(p, stripped, base_size=Pt(11))


def text_to_docx(text: str, doc_type: str = "resume") -> bytes:
    """
    Render markdown produced by the document-writer node into a clean,
    submission-ready .docx. doc_type controls layout: "resume" gets the
    name/contact header treatment, bordered section headings, and tight
    bullet spacing; "cover-letter" gets standard business-letter spacing.
    """
    doc = Document()

    for section in doc.sections:
        margin = Inches(0.75) if doc_type == "resume" else Inches(1)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10.5 if doc_type == "resume" else 11)

    lines = text.splitlines()

    if doc_type == "resume":
        _markdown_to_resume_docx(doc, lines)
    else:
        _markdown_to_cover_letter_docx(doc, lines)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
